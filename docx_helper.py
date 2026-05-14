from pathlib import Path
from docx import Document
from docx.table import _Cell, Table

def _replace_in_paragraph(paragraph, mapping):
    """Replace placeholders inside a single paragraph."""
    for run in paragraph.runs:
        text = run.text
        for key, value in mapping.items():
            if key in text:
                run.text = text.replace(key, str(value))

def _replace_in_table(table, mapping):
    """Replace placeholders inside every cell of a table."""
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell(cell, mapping)

def _replace_in_cell(cell: _Cell, mapping):
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, mapping)
    # Nested tables are rare but possible
    for tbl in cell.tables:
        _replace_in_table(tbl, mapping)

def render_template(
    template_path: str | Path,
    output_path: str | Path,
    **fields,
) -> Path:
    """
    Render a .docx template.

    Parameters
    ----------
    template_path : path to the .docx containing placeholders like {{NAME}}
    output_path   : where the generated file will be saved
    **fields      : key/value pairs used for replacement (e.g. NAME="John")

    Returns
    -------
    Path to the rendered document.
    """
    doc = Document(template_path)

    # Create mapping with the same placeholder style you used in the template
    mapping = {f"{{{{{k}}}}}": v for k, v in fields.items()}

    # Paragraphs in the main body
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    # Tables
    for tbl in doc.tables:
        _replace_in_table(tbl, mapping)

    # Header & footer (optional)
    for section in doc.sections:
        for header in (section.header, section.footer):
            for para in header.paragraphs:
                _replace_in_paragraph(para, mapping)

    doc.save(output_path)
    return Path(output_path)
